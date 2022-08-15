from ast import Try
from genericpath import isfile
from logging import exception
from multiprocessing.dummy import Array
from xmlrpc.client import Boolean
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.action_chains import ActionChains
import datetime
import time
import os
import shutil
from docx import Document
from docx.shared import Cm
from docx.enum.section import WD_ORIENT
from BaseService import BaseService  # 處理文件的直向/橫向
from PIL import Image
from selenium.common.exceptions import NoSuchElementException
import rotatescreen


class BondingShot(BaseService):

    def do_action(self):

        print(self._nowTime.strftime('%Y-%m-%d %H:%M:%S'))

        # 刪除三天前資料夾
        if self._nowTime.strftime('%H') == '00':
            _removeFolder = (
                self._nowTime + datetime.timedelta(days=-3)).strftime('%Y%m%d')
            _dirList = os.listdir(self._shareFolderPath)

            for dirName in _dirList:
                if int(dirName) <= int(_removeFolder):
                    shutil.rmtree(f'{self._shareFolderPath}\{dirName}')

        options = webdriver.ChromeOptions()
        exclude = ['enable-logging','enable-automation', 'ignore-certificate-errors'] # 停用 print log、關閉"正受到自動測試軟件的控製"訊息、忽略憑證錯誤
        options.add_experimental_option('excludeSwitches', exclude)
        # options.add_argument('--headless')
        driver = webdriver.Chrome(executable_path=".\\Driver\\chromedriver.exe", chrome_options=options)
        driver.get(self._url)

        # 計算分時看板站點圖表數量
        _chartCnt = driver.find_elements_by_xpath(
            "//img[@style='height:230px;width:480px;border-width:0px;Position:absolute;top:10px;left:5px;']")

        # 新增本日資料夾
        if(not os.path.isdir(f'{self._shareFolderPath}\{self._strDate}')):
            os.mkdir(f'{self._shareFolderPath}\{self._strDate}')

        screen = rotatescreen.get_primary_display()
        start_pos = screen.current_orientation

        if start_pos != 0:
            screen.rotate_to(0)

        # 小於7個站點, 顯示調整為直向, 截一張圖
        if len(_chartCnt) <= 7:
            self.browser_setting(driver, self._rotationZoom)
            
            pos = abs((start_pos - 90) % 360)
            screen.rotate_to(pos)

            _picName = f"{self._docFileName}_{self._pngTime}.png"
            driver.get_screenshot_as_file(_picName)

            self.move_shot_pic(_picName)

            driver.quit()

            pos = abs((pos + 90) % 360)
            screen.rotate_to(pos)

        # 大於7個站點, 橫向顯示, 分區截圖再合併
        else:
            self.browser_setting(driver, self._zoom)

            elementArray = ["Chart1", "Chart4", "Chart7",
                            "Chart" + str(len(_chartCnt) - 1)]

            for idx, ele in enumerate(elementArray):
                _picName = f"{self._docFileName}_{self._strTime}_t{idx + 1}.png"
                if((idx + 1) == 1):
                    # 將滾動條拖到最底部
                    js = F"var action=document.documentElement.scrollTop={self._scrollTop}"
                    driver.execute_script(js)
                    driver.get_screenshot_as_file(_picName)
                else:
                    try:
                        # Handle element not found
                        charts = driver.find_element_by_id(ele)
                        action = ActionChains(driver)
                        action.move_to_element(charts).perform()
                        time.sleep(1)
                        driver.get_screenshot_as_file(_picName)
                    except NoSuchElementException:
                        continue

                self._picNameArray.append(_picName)

                # time.sleep(1)
                # if(os.path.exists(f'.\{_picName}')):
                #     shutil.move(f'.\{_picName}',
                #                 f'{self._shareFolderPath}\{self._strDate}\{_picName}')

            driver.quit()

            time.sleep(1)

            self.do_spell()

            # region local test (截圖插入 word 轉成 PDF)

            '''
            if(not os.path.isdir(f'.\Report\{_strDate}')):
                os.mkdir(f'.\Report\{_strDate}')
             
            _picName = f"Bonding_Output_{_strTime}_1.png"
            _picNameArray.append(_picName)
            driver.get_screenshot_as_file(_picName)
            shutil.move(f'.\{_picName}',f'.\Report\{_strDate}\{_picName}')
            # charts = driver.find_element_by_tag_name("details")
            charts = driver.find_element_by_id("Chart8")
            action = ActionChains(driver)
            action.move_to_element(charts).perform()
            time.sleep(2)
            _picName2 = f"Bonding_Output_{_strTime}_2.png"
            _picNameArray.append(_picName2)
            driver.get_screenshot_as_file(_picName2)
            shutil.move(f'.\{_picName2}',f'.\Report\{_strDate}\{_picName2}')
            
            driver.quit()
            
            document = Document()
            section = document.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            new_width, new_height = Cm(29.7), Cm(21)
            
            section.page_width = new_width
            
            section.page_height= new_height
            section.left_margin=Cm(1.27)
            section.right_margin=Cm(1.27)
            section.top_margin=Cm(1.27)
            section.bottom_margin=Cm(1.27)
            
            p = document.add_paragraph()
            r = p.add_run()
            for picName in _picNameArray:
                r.add_picture(f'.\Report\{_strDate}\{picName}',width=Cm(27))
                
            document.save(f'.\Report\\{_strDate}\\{_docFileName}.docx')
            
            source_doc = aw.Document(f'.\Report\\{_strDate}\\{_docFileName}.docx')
            # Save as PDF
            source_doc.save(f'.\Report\\{_strDate}\\{_docFileName}.pdf')
            '''
            # endregion

    """
    設定瀏覽器參數
    """
    def browser_setting(self, initDriver: webdriver.Chrome, zoom: int) -> None:
        initDriver.get('chrome://settings/')
        initDriver.execute_script(
            F"chrome.settingsPrivate.setDefaultZoom({zoom});")
        initDriver.get(self._url)
        initDriver.fullscreen_window()

    """
    移動圖片至MApp發送資料夾
    """
    def move_shot_pic(self, pinName: str) -> None:
        try:
            time.sleep(1)
            if(os.path.exists(f'.\{pinName}')):
                shutil.move(f'.\{pinName}',
                            f'{self._shareFolderPath}\{self._strDate}\{pinName}')
        except:
            if(os.path.exists(f'.\{pinName}')):
                shutil.move(f'.\{pinName}',
                            f'{self._shareFolderPath}\{self._strDate}\{pinName}')
    
    """
    截圖數量判斷要兩兩合併或直接合併成一張
    """
    def do_spell(self):

        # 兩張以上的偶數張數, 兩兩合併為一張新圖, 暫解全部合併圖檔過大上傳模糊問題
        if len(self._picNameArray) > 2 and len(self._picNameArray) % 2 == 0:
            for i in range(0, len(self._picNameArray), 2):
                self.do_process(
                    [self._picNameArray[i], self._picNameArray[i+1]], i+1)
        else:
            self.do_process(self._picNameArray, 1)

    """
    合併截圖
    """
    def do_process(self, picArray: list, cnt: int):

        try:
            im = Image.open(F'.\{picArray[0]}')  # 開啟圖片
            xsize, ysize = im.size
            # 產生一張全黑圖片, 大小 x:長與截圖相同 y:長乘上截圖數量
            bg = Image.new('RGB', (xsize, ysize*len(picArray)), '#000000')
            im.close()

            for i, ele in enumerate(picArray):
                # 開啟圖片
                img = Image.open(F'.\{ele}')
                # 貼上圖片 左上座標(0, 0)
                bg.paste(img, (0, i*ysize))
                img.close()
            
            # 因有機率會因權限問題會失敗, 故再重試一次
            try:
                bg.save(
                    F'{self._shareFolderPath}\{self._strDate}\{self._docFileName}_{self._pngTime}_{str(cnt)}.png')
            except:
                bg.save(
                    F'{self._shareFolderPath}\{self._strDate}\{self._docFileName}_{self._pngTime}_{str(cnt)}.png')

            # 刪除已合併截圖
            for tmp in picArray:
                if os.path.isfile(F'.\{tmp}'):
                    os.remove(F'.\{tmp}')

        except Exception as ex:
            print(str(ex))
