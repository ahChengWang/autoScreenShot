from selenium import webdriver
from selenium.webdriver.chrome.service import Service
# from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.action_chains import ActionChains
from datetime import datetime
import os
import shutil
import time
from docx import Document
from docx.shared import Cm
from docx.enum.section import WD_ORIENT #處理文件的直向/橫向
import requests
import base64
from datetime import date
import urllib.parse
import aspose.words as aw
import glob
import json


options = webdriver.ChromeOptions()
options.add_experimental_option('excludeSwitches', ['enable-logging'])
driver = webdriver.Chrome(executable_path=".\\Driver\\chromedriver.exe", chrome_options=options)
# driver = webdriver.Chrome(executable_path=".\\chromedriver_win32_102.0.5005.61\\chromedriver.exe", chrome_options=options)
driver.get('chrome://settings/')
driver.execute_script('chrome.settingsPrivate.setDefaultZoom(0.88);')
driver.get("http://10.132.133.164/Function/ENG1/Bonding_Output.aspx") # 3F Bonding

# driver.get("https://www.edh.tw/article/30906") # 測試

driver.fullscreen_window()

_strDate = datetime.now().strftime('%Y%m%d')
_strTime = datetime.now().strftime('%y%m%d%H%M%S')
_pngTime = datetime.now().strftime('%y%m%d%H')
_docFileName = f'Bonding_Report_{_strTime}'
_picNameArray = []
_shareFolderPath = f'Z:\\02-共用資料區(3G)\\03-戰情\\3F_Bonding\\Daily'

time.sleep(1)

if(not os.path.isdir(f'{_shareFolderPath}\{_strDate}')):
    os.mkdir(f'{_shareFolderPath}\{_strDate}')

for i in range(1,4):
    _picName = f"Bonding_Output_{_strTime}_{i}.png"
    _picNameArray.append(_picName)
    if(i==1):
        driver.get_screenshot_as_file(_picName)
    else:
        charts = driver.find_element_by_id("Chart5" if i == 2 else "Chart8")
        action = ActionChains(driver)
        action.move_to_element(charts).perform()
        time.sleep(2)
        driver.get_screenshot_as_file(_picName)
    if(os.path.exists(f'.\{_picName}')):
        shutil.move(f'.\{_picName}',f'{_shareFolderPath}\{_strDate}\{_picName}')

driver.quit()


# region local test (截圖插入 word 轉成 PDF)
'''

if(not os.path.isdir(f'.\Report\{_strDate}')):
    os.mkdir(f'.\Report\{_strDate}')
 
_picName = f"Bonding_Output_{_strTime}_1.png"
_picNameArray.append(_picName)
driver.get_screenshot_as_file(_picName)
shutil.move(f'.\{_picName}',f'.\Report\{_strDate}\{_picName}')
charts = driver.find_element_by_tag_name("details")
action = ActionChains(driver)
action.move_to_element(charts).perform()
time.sleep(2)
_picName2 = f"Bonding_Output_{_strTime}_2.png"
_picNameArray.append(_picName2)
driver.get_screenshot_as_file(_picName2)
shutil.move(f'.\{_picName2}',f'.\Report\{_strDate}\{_picName2}')

driver.quit()

document = Document()
section = document.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
new_width, new_height = Cm(29.7), Cm(21)

section.page_width = new_width

section.page_height= new_height
section.left_margin=Cm(1.27)
section.right_margin=Cm(1.27)
section.top_margin=Cm(1.27)
section.bottom_margin=Cm(1.27)

p = document.add_paragraph()
r = p.add_run()
for picName in _picNameArray:
    r.add_picture(f'.\Report\{_strDate}\{picName}',width=Cm(27))
    
document.save(f'.\Report\\{_strDate}\\{_docFileName}.docx')

source_doc = aw.Document(f'.\Report\\{_strDate}\\{_docFileName}.docx')
# Save as PDF
source_doc.save(f'.\Report\\{_strDate}\\{_docFileName}.pdf')
'''

# endregion 

